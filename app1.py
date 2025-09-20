from flask import Flask, render_template, request, send_file, session, redirect, url_for, flash
import sqlite3
import os
import tempfile
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI

from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# ----------------- Config -----------------
DB_PATH = "companies.db"
APP_SECRET = os.getenv("FLASK_SECRET", "dev-secret")
HIDE_VSP_FROM_DOCX = False          # Hide VSP in Word export
VSP_VISIBLE_IN_UI = True           # Show VSP in result page
OPENAI_MODEL = "gpt-4o"
TEMPERATURE_VSP = 0.9
MAX_TOKENS_VSP = 1700
TEMPERATURE_EXEC = 0.9
MAX_TOKENS_EXEC = 1700

# ----------------- Flask -----------------
app = Flask(__name__)
app.secret_key = APP_SECRET

# ----------------- OpenAI -----------------
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    organization=os.getenv("OPENAI_ORG_ID")
)

# ----------------- In-memory storage -----------------
storage = {}

# ----------------- DB Helpers -----------------
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS companies (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            industry TEXT,
            services TEXT,
            differentiators TEXT,
            contact_email TEXT,
            contact_phone TEXT,
            website TEXT,
            notes TEXT
        )
    """)
    conn.commit()
    conn.close()

def insert_company(profile):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute(
        """INSERT INTO companies 
        (name, industry, services, differentiators, contact_email, contact_phone, website, notes) 
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
        (profile.get("name"), profile.get("industry"), profile.get("services"),
         profile.get("differentiators"), profile.get("contact_email"), profile.get("contact_phone"),
         profile.get("website"), profile.get("notes"))
    )
    conn.commit()
    company_id = cur.lastrowid
    conn.close()
    return company_id

def get_companies():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT id, name, industry FROM companies ORDER BY name")
    rows = cur.fetchall()
    conn.close()
    return rows

def get_company(company_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""SELECT id, name, industry, services, differentiators, contact_email, 
                   contact_phone, website, notes 
                   FROM companies WHERE id = ?""", (company_id,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return None
    keys = ["id", "name", "industry", "services", "differentiators", 
            "contact_email", "contact_phone", "website", "notes"]
    return dict(zip(keys, row))

# ----------------- Text cleaning -----------------
def clean_text_block(text: str) -> str:
    if not text:
        return ""
    text = str(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"•\s*", "- ", text, flags=re.MULTILINE)
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()

# ----------------- Prompts -----------------

# Single-pass Executive Summary prompt, with stronger rules only for:
# - Solution Overview
# - How We Will Deliver
# - Why {provider_name}
EXEC_PROMPT_TEMPLATE = """
You are a senior management consultant. Using the provider profile and the VSP, produce a polished,
client-ready Executive Summary in well formatted plain text (no Markdown, no ##, no **).
Tone must mirror the business-driven, persuasive style of the VSP — sharp, opportunity-focused, and executive-level.

Requirements:
- Length: 600–900 words.
- Use EXACTLY these headings once, in order:
  1) Introduction
  2) Our Understanding of Your Goals
  3) Our Approach to Meeting Your Goals
  4) Solution Overview
  5) How We Will Deliver
  6) Why {provider_name}
  7) Closing Call-to-Action
- Headings must appear exactly once and in order. Do not repeat any heading.
- After “Our Understanding of Your Goals”, maintain forward flow: Approach → Solution → Delivery → Why.
- Frame the client positively (readiness/opportunity). Avoid weakness/problem language.
- Use "-" for bullets. No other symbols. No Markdown. No placeholders.
- Do not invent facts. Reuse exact phrases from the VSP; especially for sections 4–6.
- Adapt tone, vocabulary, and emphasis according to the Recipient Role specified in the client context (the cahnges from role to role to be subtle without losing the essence of what an executive summary is ):
  - For a CEO: emphasize innovation, long-term strategic advantage, and market leadership.
  - For a CFO: emphasize ROI, cost savings, margin improvement, EBITDA impact, and financial resilience.
  - For a CIO/CTO: emphasize technical robustness, scalability, integration, compliance, and innovation in IT systems.
  - For a Head of Sales / CMO: emphasize business value, revenue growth, customer experience, and competitive differentiation.
  - For an Operations Director: emphasize efficiency, risk mitigation, governance, and process excellence.
  - For other roles: infer focus areas logically while maintaining clarity and professionalism.


Section specifics:
- Our Approach to Meeting Your Goals:
  - 2 lines descriptive statement followed by explanable 3–4 bullets.
  - Explain your methodology, value engineering levers (e.g., EBITDA, margins, cost-to-serve, working capital, time-to-value), and change thesis.
  - Map client goals → approach elements → measurable business outcomes. Impress with provider capability.
- Solution Overview:
  - 3–5 bullets. EVERY bullet must reuse at least one exact phrase from the VSP “Proposed Solution”.
  - Each bullet must map the module to an explicit business outcome (efficiency, EBITDA, margins, costs, working capital, patient experience).
- How We Will Deliver:
  - 3–5 bullets. Focus on execution mechanics: governance cadence, risk mitigation, phased rollout, joint ownership/BOT, enablement/training.
  - Tie each bullet to measurement and confidence (baselines, KPIs, ranges ok).
- Why {provider_name}:
  - 3–5 bullets. Reuse differentiators from provider profile and VSP (e.g., certifications, domain expertise, security-first, references).
  - Each bullet must explicitly state the client value of that differentiator.

- Closing Call-to-Action:
  - 2–3 sentences. Formal, decisive, non-casual.
  - Invite next-step meeting (e.g., governance/kickoff/strategy session) AND include the provider’s contact email and phone from the provider profile if present.

Inputs:
- PROVIDER_PROFILE:
{provider_profile}
- VSP:
{vsp_resp}
- Website of provider:
{website}
- CLIENT_CONTEXT:
{client_context}
"""



def get_vsp_prompt(company, provider_profile_text, client_context):
    return f"""
You are a senior management consultant. Based on the provider profile and client context,
generate a Value Selling Points (VSP) document.

Guidelines:
- Plain text only (no Markdown, no symbols).
- Each bullet must be a strong business phrase (1–2 lines).
- Structure exactly:

Case for Change
- ...
Business Value for the Client
- ...
{company['name']} Proposed Solution
- ...

Inputs:
PROVIDER_PROFILE:
{provider_profile_text}

CLIENT_CONTEXT:
{client_context}
"""

# ----------------- Routes -----------------
@app.route("/setup", methods=["GET", "POST"])
def setup():
    if request.method == "POST":
        profile = {
            "name": request.form.get("name", "").strip(),
            "industry": request.form.get("industry", "").strip(),
            "services": request.form.get("services", "").strip(),
            "differentiators": request.form.get("differentiators", "").strip(),
            "contact_email": request.form.get("contact_email", "").strip(),
            "contact_phone": request.form.get("contact_phone", "").strip(),
            "website": request.form.get("website", "").strip(),
            "notes": request.form.get("notes", "").strip()
        }
        if not profile["name"]:
            flash("Company name is required", "warning")
            return render_template("setup.html", profile=profile)
        company_id = insert_company(profile)
        flash(f"Company profile '{profile['name']}' created.", "success")
        return redirect(url_for("index", company_id=company_id))
    return render_template("setup.html")

@app.route("/", methods=["GET", "POST"])
def index():
    companies = get_companies()
    default_company_id = request.args.get("company_id")

    if request.method == "POST":
        # ---------- Mandatory field checks ----------
        errors = []
        try:
            company_id = int(request.form.get("company_id"))
        except Exception:
            company_id = None
        if not company_id:
            errors.append("Please select a provider.")

        client_name = request.form.get("client_name", "").strip()
        if not client_name:
            errors.append("Client Name is required.")
        client_industry = request.form.get("client_industry", "").strip()
        if not client_industry:
            errors.append("Client Industry is required.")
        client_goals = request.form.get("client_goals", "").strip()
        if not client_goals:
            errors.append("Client Goals / Challenges are required.")
        proposal_modules = request.form.get("proposal_modules", "").strip()
        if not proposal_modules:
            errors.append("Proposed Solutions / Modules are required.")
        recipient_role = request.form.get("recipient_role", "").strip()
        if not recipient_role:
         errors.append("Recipient Role is required.")
        execution_model = request.form.get("execution_model", "").strip()
        extra_notes = request.form.get("extra_notes", "").strip()
        

        if errors:
            for e in errors:
                flash(e, "warning")
            return render_template("index.html", companies=companies, default_company_id=default_company_id)

        # ---------- Get company ----------
        company = get_company(company_id)
        if not company:
            flash("Selected company not found.", "danger")
            return redirect(url_for("setup"))

        # ---------- Build client context ----------
        client_context = f"""
Client Name: {client_name}
Client Industry: {client_industry}
Goals/Challenges: {client_goals}
Proposed Modules: {proposal_modules}
Recipient Role: {recipient_role}
Execution Model: {execution_model}
Additional Notes: {extra_notes}
""".strip()

        # ---------- Prepare session ----------
        session_id = os.urandom(8).hex()
        session["id"] = session_id
        storage[session_id] = {"client_context": client_context, "draft": "", "vsp": "", "company_id": company_id}

        provider_profile_text = f"{company['name']} (Industry: {company['industry']}). Services: {company['services']}. Differentiators: {company['differentiators']}. Website: {company['website']}. Contact: {company['contact_email']} | {company['contact_phone']}."

        # ---------- Generate VSP ----------
        vsp_prompt = get_vsp_prompt(company, provider_profile_text, client_context)
        try:
            vsp_resp = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role": "system", "content": "You are an expert proposal writer."},
                    {"role": "user", "content": vsp_prompt}
                ],
                temperature=TEMPERATURE_VSP,
                max_completion_tokens=MAX_TOKENS_VSP
            )
            vsp_text = clean_text_block(vsp_resp.choices[0].message.content.strip())
        except Exception as e:
            vsp_text = f"VSP generation failed: {e}"

        storage[session_id]["vsp"] = vsp_text

        # ---------- Generate Executive Summary (single pass) ----------
        exec_prompt = EXEC_PROMPT_TEMPLATE.format(
            provider_profile=provider_profile_text,
            client_context=client_context,
            provider_name=company['name'],
            website=company["website"],
            vsp_resp=vsp_text
        )

        try:
            exec_resp = client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role": "system", "content": "You are a senior management consultant writing client-ready executive summaries."},
                    {"role": "user", "content": exec_prompt}
                ],
                temperature=TEMPERATURE_EXEC,
                max_completion_tokens=MAX_TOKENS_EXEC
            )
            exec_text = clean_text_block(exec_resp.choices[0].message.content.strip())
        except Exception as e:
            exec_text = f"Executive Summary generation failed: {e}"

        # ---------- Replace any model-written closing with formal CTA ----------
        prov = company.get('name', 'Provider')
        client_disp = client_name if client_name else "the client"
        formal_cta = (
            f"{prov} recommends moving forward with a phased engagement to realize measurable operational efficiencies within the first year. "
            f"We are prepared to initiate governance reviews, align executive stakeholders, and formalize next steps to ensure {client_disp} achieves "
            f"sustainable improvements in patient satisfaction, cost efficiency, and compliance readiness."
        )
        # Remove any existing closing section the model might have written

        storage[session_id]["draft"] = exec_text

        return redirect(url_for("result"))

    # ---------- GET request ----------
    return render_template("index.html", companies=companies, default_company_id=default_company_id)

# ----------------- Result Route -----------------
@app.route("/result", methods=["GET", "POST"])
def result():
    session_id = session.get("id")
    if not session_id or session_id not in storage:
        flash("No active proposal found. Please generate one.", "warning")
        return redirect(url_for("index"))

    data = storage[session_id]
    draft = data.get("draft", "")
    vsp = data.get("vsp", "")
    client_context = data.get("client_context", "")
    company_id = data.get("company_id")
    company = get_company(company_id) if company_id else None

    if request.method == "POST":
        # ---------- Refine ----------
        if "refine" in request.form:
            refine_input = request.form.get("refine_prompt", "").strip()
            if not refine_input:
                flash("Refine instructions cannot be empty.", "warning")
                return render_template("result.html", draft=draft, vsp=vsp,
                                       context=client_context, company=company, show_vsp=VSP_VISIBLE_IN_UI)

            refine_prompt = f"""
Refine the Executive Summary below using these instructions exactly:
{refine_input}

Executive Summary:
{draft}

Rules:
- Keep section order intact (Introduction → … → Closing Call-to-Action).
- Reuse VSP phrases in Solution Overview, How We Will Deliver, and Why {company['name'] if company else 'Provider'}.
- Use "-" for bullets where bullets already exist.
- Do NOT add Markdown or placeholders.
- Preserve the existing Closing Call-to-Action format and contact details.
"""
            try:
                refine_resp = client.chat.completions.create(
                    model=OPENAI_MODEL,
                    messages=[
                        {"role": "system", "content": "You are a professional consultant refining executive summaries."},
                        {"role": "user", "content": refine_prompt}
                    ],
                    temperature=TEMPERATURE_EXEC,
                    max_completion_tokens=MAX_TOKENS_EXEC
                )
                draft = clean_text_block(refine_resp.choices[0].message.content.strip())

                # Re-append formal CTA to be safe
                prov = company.get('name', 'Provider') if company else "Provider"
                client_disp = "the client"
                formal_cta = (
                    f"{prov} recommends moving forward with a phased engagement to realize measurable operational efficiencies within the first year. "
                    f"We are prepared to initiate governance reviews, align executive stakeholders, and formalize next steps to ensure {client_disp} achieves "
                    f"sustainable improvements in patient satisfaction, cost efficiency, and compliance readiness."
                )
                draft = re.sub(r"\n?7\)\s*Closing.*|^Closing.*", "", draft, flags=re.IGNORECASE | re.DOTALL).strip()
                draft += "\n\nClosing Call-to-Action\n" + formal_cta

                storage[session_id]["draft"] = draft
            except Exception as e:
                flash(f"Refine failed: {e}", "danger")

        # ---------- Download ----------
        elif "download" in request.form:
            doc = Document()
            doc.add_heading(f"Executive Summary by {company['name'] if company else 'Provider'}", level=0)

            # Write paragraphs with simple heading detection
            for line in draft.split("\n"):
                if not line.strip():
                    continue
                stripped = line.strip()
                lower = stripped.lower()

                if lower.startswith("introduction"):
                    title = "Introduction"
                elif lower.startswith("our understanding"):
                    title = "Our Understanding of Your Goals"
                elif lower.startswith("our approach"):
                    title = "Our Approach to Meeting Your Goals"
                elif lower.startswith("solution overview"):
                    title = "Solution Overview"
                elif lower.startswith("how we will deliver"):
                    title = "How We Will Deliver"
                elif lower.startswith("why"):
                    title = f"Why {company['name'] if company else 'Provider'}"
                elif lower.startswith("closing"):
                    title = "Closing Call-to-Action"
                else:
                    title = None

                if title:
                    para = doc.add_paragraph(title)
                    run = para.runs[0]
                    run.bold = True
                    run.font.size = Pt(14)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.space_after = Pt(10)
                elif stripped.startswith("- "):
                    para = doc.add_paragraph(stripped[2:], style="List Bullet")
                    para.paragraph_format.space_after = Pt(4)
                else:
                    para = doc.add_paragraph(stripped)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.space_after = Pt(8)

            # Contact info
            doc.add_heading("Contact Information", level=1)
            if company:
                contact_block = [
                    f"Email: {company.get('contact_email', 'N/A')}",
                    f"Phone: {company.get('contact_phone', 'N/A')}",
                    f"Website: {company.get('website', 'N/A')}"
                ]
                for item in contact_block:
                    para = doc.add_paragraph(item)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    para.paragraph_format.space_after = Pt(4)

            # Optional VSP appendix
            if not HIDE_VSP_FROM_DOCX and vsp.strip():
                doc.add_page_break()
                doc.add_heading(f"Value Selling Points by {company['name'] if company else 'Provider'}", level=0)
                for line in vsp.split("\n"):
                    if not line.strip():
                        continue
                    stripped = line.strip()
                    lower = stripped.lower()
                    if lower.startswith("case for change"):
                        title = "Case for Change"
                    elif lower.startswith("business value"):
                        title = "Business Value for the Client"
                    elif (company and lower.startswith(company['name'].lower())) or lower.startswith("proposed solution"):
                        title = f"{company['name']} Proposed Solution" if company else "Proposed Solution"
                    else:
                        title = None

                    if title:
                        para = doc.add_paragraph(title)
                        run = para.runs[0]
                        run.bold = True
                        run.font.size = Pt(14)
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.paragraph_format.space_after = Pt(10)
                    elif stripped.startswith("- "):
                        para = doc.add_paragraph(stripped[2:], style="List Bullet")
                        para.paragraph_format.space_after = Pt(4)
                    else:
                        para = doc.add_paragraph(stripped)
                        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        para.paragraph_format.space_after = Pt(6)
                
                # -------- Client Context Section --------
            if client_context:
                 doc.add_page_break()
                 doc.add_heading("Client Context", level=0)
                 for line in client_context.split("\n"):
                  if not line.strip():
                   continue
                  para = doc.add_paragraph(line.strip())
                  para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                  para.paragraph_format.space_after = Pt(6)

            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(tmp_file.name)
            return send_file(tmp_file.name, as_attachment=True, download_name="Executive_Summary.docx")

        # ---------- Finish ----------
        elif "finish" in request.form:
            storage.pop(session_id, None)
            session.clear()
            flash("Session cleared.", "info")
            return redirect(url_for("index"))

    return render_template("result.html", draft=draft, vsp=vsp,
                           context=client_context, company=company, show_vsp=VSP_VISIBLE_IN_UI)
    # Build messages list for chat UI
    

    


if __name__ == "__main__":
    init_db()
    app.run(debug=True, host="0.0.0.0", port=5000)
